# -*- coding: utf-8 -*-
import os, sys, cv2, csv, time, json, random, gc, re, ast
import threading
import concurrent.futures
import numpy as np
import fitz
import pytesseract
import openpyxl
from openpyxl import Workbook
from PIL import Image
import google.generativeai as genai

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import xlrd
except ImportError:
    xlrd = None

from utils import (
    auto_adjust_excel_column_width,
    analyze_column_profile,
    get_profile_similarity,
    parse_row_data,
    apply_text_inheritance,
    merge_2d_arrays_horizontally,
    sanitize_excel_text
)

# ==============================
# AI抽出・データ集約タスク
# ==============================
def check_tesseract_installation():
    if sys.platform.startswith("win"):
        tess_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(tess_path): pytesseract.pytesseract.tesseract_cmd = tess_path
    try: pytesseract.get_tesseract_version()
    except Exception: raise Exception("Tesseract OCRがインストールされていないか、PATHが通っていません。")

def extract_tesseract_task(files, save_dir, options, ui):
    check_tesseract_installation()
    files = [f for f in files if f.lower().endswith(".pdf")]
    if not files: raise Exception("PDFファイルが含まれていません。")
    out_format = options.get("out_format", "xlsx")
    crop_regions = options.get("crop_regions", [])
    
    for i, f in enumerate(files, 1):
        if ui.is_cancelled(): return
        ui.update_overall(i, len(files), f"全体の進捗 ( {i} / {len(files)} ファイル )")
        base = os.path.splitext(os.path.basename(f))[0]
        doc = fitz.open(f)
        digits = max(2, len(str(len(doc))))
        total_pages = len(doc)
        for page_num in range(total_pages):
            if ui.is_cancelled(): return
            ui.set_indeterminate(f"Tesseractで解析中... ( {page_num+1} / {total_pages} ページ )")
            pix = doc[page_num].get_pixmap(dpi=300)
            img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
            if pix.n == 4: img_array = cv2.cvtColor(img_array, cv2.COLOR_RGBA2RGB)
            elif pix.n == 1: img_array = cv2.cvtColor(img_array, cv2.COLOR_GRAY2RGB)
            
            cropped_images = []
            if crop_regions:
                h, w = img_array.shape[:2]
                for (rx1, ry1, rx2, ry2) in crop_regions:
                    cropped_images.append(img_array[int(min(ry1, ry2)*h):int(max(ry1, ry2)*h), int(min(rx1, rx2)*w):int(max(rx1, rx2)*w)])
            else: cropped_images.append(img_array)
                
            all_regions_data = []
            for crop_img in cropped_images:
                try:
                    text = pytesseract.image_to_string(Image.fromarray(crop_img), lang="jpn+eng")
                    lines = [line.strip() for line in text.split('\n') if line.strip()]
                    if lines:
                        if (crop_regions and crop_img.shape[0] > crop_img.shape[1] * 1.5) or (len(lines) > 1 and all(len(l) <= 2 for l in lines)):
                            all_regions_data.append([["".join(lines)]])
                        else: all_regions_data.append([[l] for l in lines])
                    else: all_regions_data.append([[""]])
                except Exception as e: raise Exception(f"Tesseract OCRエラー: {e}")
                    
            merged_data = merge_2d_arrays_horizontally(all_regions_data)
            final_data = []
            page_info = f"{page_num+1}/{len(doc)}"
            if crop_regions:
                final_data.append(["ページ番号"] + [f"抽出範囲{idx+1}" for idx in range(len(cropped_images))])
            else:
                final_data.append(["ページ番号", "抽出テキスト"])
            for row in merged_data: final_data.append([page_info] + row)
            
            save_path = os.path.join(save_dir, f"{base}_Page_{str(page_num+1).zfill(digits)}_Tesseract抽出")
            if out_format == "xlsx":
                wb = Workbook(); ws = wb.active; ws.title = f"Page_{str(page_num+1).zfill(digits)}"
                for r_idx, row_data in enumerate(final_data, 1):
                    for c_idx, val in enumerate(row_data, 1): 
                        ws.cell(row=r_idx, column=c_idx, value=sanitize_excel_text(val))
                auto_adjust_excel_column_width(ws); wb.save(f"{save_path}.xlsx")
            elif out_format == "csv":
                with open(f"{save_path}.csv", "w", encoding="utf-8-sig", newline="") as f_out: csv.writer(f_out).writerows(final_data)
            elif out_format == "txt":
                with open(f"{save_path}.txt", "w", encoding="utf-8") as f_out:
                    for row_data in final_data: f_out.write("\t".join(row_data) + "\n")
            elif out_format == "json":
                with open(f"{save_path}.json", "w", encoding="utf-8") as f_out: json.dump(final_data, f_out, ensure_ascii=False, indent=2)
            elif out_format == "md":
                with open(f"{save_path}.md", "w", encoding="utf-8") as f_out:
                    if final_data:
                        f_out.write("| " + " | ".join(map(str, final_data[0])) + " |\n")
                        f_out.write("|" + "|".join(["---"] * len(final_data[0])) + "|\n")
                        for row in final_data[1:]: f_out.write("| " + " | ".join(map(lambda x: str(x).replace('\n', '<br>'), row)) + " |\n")
            elif out_format == "docx":
                if Document is None: raise Exception("python-docx ライブラリがインストールされていません。")
                doc_out = Document()
                if final_data:
                    table = doc_out.add_table(rows=len(final_data), cols=max(len(r) for r in final_data))
                    table.style = 'Table Grid'
                    for r_idx, row_data in enumerate(final_data):
                        row_cells = table.rows[r_idx].cells
                        for c_idx, val in enumerate(row_data):
                            if c_idx < len(row_cells):
                                row_cells[c_idx].text = str(val)
                doc_out.save(f"{save_path}.docx")
        doc.close(); gc.collect()
        ui.set_determinate(total_pages, total_pages, "完了")

def extract_gemini_task(files, save_dir, options, ui):
    files = [f for f in files if f.lower().endswith(".pdf")]
    if not files: raise Exception("PDFファイルが含まれていません。")
    api_key = options.get("api_key", "")
    genai.configure(api_key=api_key)
    
    models_to_try = options.get("models_to_try", ["gemini-2.5-flash"])
    out_format = options.get("out_format", "xlsx")
    crop_regions = options.get("crop_regions", [])

    if out_format in ["csv", "xlsx", "json", "md", "docx"]:
        if crop_regions:
            prompt = """
            あなたは優秀なデータ入力オペレーターです。添付された複数の画像（抽出領域）のテキストを順番に読み取り、画像ごとに分割したJSONデータを作成してください。
            
            【特別ルール（絶対厳守）】
            - 画像は複数枚渡されます。添付された順番通りに、1枚目を「領域1」、2枚目を「領域2」として処理してください。
            - 1つの画像（領域）につき、出力は「1つの配列（列）」にまとめます。領域内での列の分割（セル分け）は絶対にしないでください。
            - 1行分のデータは、空白などで区切られていても分割せず、すべて1つの文字列としてまとめてください。
            - 縦書きのテキスト（文字が縦に並んでいるもの）は、1文字ずつ分割したり改行したりせず、必ず繋げて横書きの1つの文字列に変換してください。

            【出力形式（絶対厳守）】
            以下のように、画像の枚数と同じ長さの配列（リストのリスト）を持つJSON形式で出力してください。
            {
              "regions": [
                [ "画像1の1行目...", "画像1の2行目..." ],
                [ "画像2の1行目...", "画像2の2行目..." ]
              ]
            }
            """
        else:
            prompt = """
            あなたは優秀なデータ入力オペレーターです。添付された図面管理台帳などの表画像を読み取り、正確なJSONデータを作成してください。
            
            【データの分離精度を最優先する特別ルール（超重要）】
            - データの見た目や文脈の意味よりも、「表の縦の罫線」や「文字列間の大きな空白」などの【物理的な列の区切り】を絶対的な基準として最優先し、物理的に分かれているデータは必ず別の要素として分割してください。
            - 意味的に繋がっているように見えても、罫線や空白で区切られていれば絶対に1つの要素に結合しないでください。
            - データが存在しない「空白セル」の場合は無視せず、必ず `""` (空文字) を該当位置に挿入し、列の位置を厳密に揃えてください。
            - 1行のデータを丸ごと1つの文字列に繋げて出力することは絶対に禁止します。必ず各セルごとにリスト内で分割してください。
            - 縦書きのテキスト（文字が縦に並んでいるもの）は、1文字ずつ分割したり複数行に分けたりせず、必ず繋げて横書きの1つの文字列に変換して出力してください。

            【出力形式（絶対厳守）】
            シンプルな配列（リスト）で出力してください。
            {
              "header": ["列1名", "列2名", "列3名", ...],
              "rows": [
                ["データ1", "データ2", "", "データ4", ...],
                ["データ1", "データ2", "データ3", "", ...]
              ]
            }
            """
        generation_config = {"response_mime_type": "application/json"}
    else:
        if crop_regions:
            prompt = """
            添付された複数の画像に記載されているテキストを順番に読み取り、プレーンテキストとして出力してください。
            各画像（領域）のテキストの間には、必ず「===REGION_SPLIT===」という区切り文字を挿入してください。
            縦書きの文章は改行を取り除き、横書きに変換して出力してください。
            """
        else:
            prompt = "この画像に記載されている手書きの文字や文章を可能な限り正確に読み取り、プレーンテキストとして出力してください。また、縦書きの文章は改行を取り除き、横書きに変換して出力してください。"
        generation_config = None

    api_plan = options.get("api_plan", "free")
    api_rpm = options.get("api_rpm", 12)
    MAX_RPM = api_rpm
    
    if MAX_RPM <= 0:
        MAX_RPM = 1
        
    MIN_REQUEST_INTERVAL = (60.0 / MAX_RPM) * 1.05 # 確実な間隔を空けるための5%マージン

    request_timestamps = []
    last_request_time = [0.0]
    rate_limit_lock = threading.Lock()
    
    if api_plan == "free":
        max_workers = 1           # 完全に直列処理として実行
        BATCH_SIZE = 1            # 1ページずつ確実に処理
    else:
        # 課金枠: 指定されたRPMに応じて並列度を調整 (上限10並列)
        max_workers = min(10, max(1, MAX_RPM // 30))
        BATCH_SIZE = min(10, max_workers)

    def process_api_request_for_page(imgs, page_num):
        max_retries = 8
        extracted_text, success, last_error = "", False, ""
        num_regions = len(imgs)
        
        for attempt in range(max_retries):
            if ui.is_cancelled(): return "", False, "ユーザーによる中止"
            required_sleep = 0.0
            is_fatal_quota = False
            
            for model_name in models_to_try:
                if ui.is_cancelled(): return "", False, "ユーザーによる中止"
                wait_t = 0.0
                rpm_wait_t = 0.0
                
                with rate_limit_lock:
                    current_time = time.time()
                    elapsed = current_time - last_request_time[0]
                    
                    if last_request_time[0] > 0 and elapsed < MIN_REQUEST_INTERVAL:
                        wait_t = MIN_REQUEST_INTERVAL - elapsed
                        
                    expected_run_time = current_time + wait_t
                    
                    nonlocal request_timestamps
                    request_timestamps = [t for t in request_timestamps if expected_run_time - t < 60.0]
                    
                    if len(request_timestamps) >= MAX_RPM:
                        rpm_wait_t = 60.0 - (expected_run_time - request_timestamps[0])
                        if rpm_wait_t > 0:
                            expected_run_time += rpm_wait_t
                            
                    last_request_time[0] = expected_run_time
                    request_timestamps.append(expected_run_time)

                total_wait = wait_t + rpm_wait_t
                if total_wait > 0:
                    ui.set_indeterminate(f"通信間隔調整中... (約 {total_wait:.1f}秒)")
                    wait_step = 0.5
                    while total_wait > 0:
                        if ui.is_cancelled(): return "", False, "ユーザーによる中止"
                        time.sleep(min(wait_step, total_wait))
                        total_wait -= wait_step
                
                if num_regions > 1:
                    ui.set_indeterminate(f"AI解析中... ( P.{page_num+1} | {num_regions}箇所の領域を一括処理 )")
                else:
                    ui.set_indeterminate(f"AI解析中... ( P.{page_num+1} )")
                
                try:
                    model = genai.GenerativeModel(model_name)
                    contents = [prompt] + imgs
                    if generation_config:
                        response = model.generate_content(contents, generation_config=generation_config)
                    else:
                        response = model.generate_content(contents)
                        
                    if not response.parts: raise Exception("安全フィルタ等によりブロックされました。")
                    extracted_text = response.text.strip()
                    success = True
                    break
                except Exception as api_err: 
                    err_str = str(api_err)
                    if "429" in err_str or "Quota" in err_str:
                        m = re.search(r'retry in ([\d\.]+)s', err_str)
                        if not m: m = re.search(r'seconds:\s*(\d+)', err_str)
                        if m:
                            required_sleep = float(m.group(1)) + 2.0 # 猶予を持たせる
                            last_error = f"API制限(429): Google側の制限で約{int(required_sleep)}秒待機します。"
                        else:
                            if "perday" in err_str.lower():
                                last_error = f"API制限(1日の上限): {model_name}の1日の無料枠を使い切った可能性があります。"
                                is_fatal_quota = True
                            else:
                                last_error = f"API制限(429): {model_name}の利用枠上限に達しました(バースト制限)。"
                                required_sleep = 15.0 # 短いバーストの場合は長めに待機
                    elif "404" in err_str: 
                        last_error = f"モデル未発見(404): {model_name}が利用できません。"
                    else: 
                        last_error = err_str
                    continue
            
            if success: break
            if is_fatal_quota: break
            if ui.is_cancelled(): return "", False, "ユーザーによる中止"
            
            if required_sleep > 0:
                sleep_time = required_sleep + random.uniform(1.0, 2.0)
            else:
                base_sleep = min(60.0, 4.0 * (2 ** attempt))
                sleep_time = base_sleep + random.uniform(1.0, 3.0)
                
            ui.set_indeterminate(f"API制限回避のため待機中... (約 {int(sleep_time)} 秒待機)")
            wait_step = 0.5
            while sleep_time > 0:
                if ui.is_cancelled(): return "", False, "ユーザーによる中止"
                time.sleep(min(wait_step, sleep_time))
                sleep_time -= wait_step

        return extracted_text, success, last_error

    for i, f in enumerate(files, 1):
        if ui.is_cancelled(): return
        ui.update_overall(i, len(files), f"全体の進捗 ( {i} / {len(files)} ファイル )")
        base = os.path.splitext(os.path.basename(f))[0]
        doc = fitz.open(f)
        digits = max(2, len(str(len(doc))))
        total_pages = len(doc)
        
        for batch_start in range(0, total_pages, BATCH_SIZE):
            if ui.is_cancelled(): return
            batch_end = min(batch_start + BATCH_SIZE, total_pages)
            
            page_tasks = {}
            for page_num in range(batch_start, batch_end):
                pix = doc[page_num].get_pixmap(dpi=300)
                img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
                if pix.n == 4: img_array = cv2.cvtColor(img_array, cv2.COLOR_RGBA2RGB)
                elif pix.n == 1: img_array = cv2.cvtColor(img_array, cv2.COLOR_GRAY2RGB)

                cropped_images = []
                if crop_regions:
                    h, w = img_array.shape[:2]
                    for (rx1, ry1, rx2, ry2) in crop_regions:
                        cropped_images.append(img_array[int(min(ry1, ry2)*h):int(max(ry1, ry2)*h), int(min(rx1, rx2)*w):int(max(rx1, rx2)*w)])
                else: cropped_images.append(img_array)
                
                page_tasks[page_num] = {}
                for region_idx, crop_img_array in enumerate(cropped_images):
                    gray = cv2.cvtColor(crop_img_array, cv2.COLOR_RGB2GRAY)
                    blurred = cv2.GaussianBlur(gray, (3, 3), 0)
                    binary = cv2.adaptiveThreshold(blurred, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 15, 5)
                    clean_bg = np.where(binary == 255, 255, gray)
                    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                    enhanced = clahe.apply(clean_bg.astype(np.uint8))
                    blur_for_sharp = cv2.GaussianBlur(enhanced, (0, 0), 2)
                    sharp = cv2.addWeighted(enhanced, 1.5, blur_for_sharp, -0.5, 0)
                    img = Image.fromarray(cv2.cvtColor(sharp, cv2.COLOR_GRAY2RGB))
                    
                    page_tasks[page_num][region_idx] = (img, crop_img_array)
            
            page_results = {}
            with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
                futures = {}
                for page_num, regions in page_tasks.items():
                    imgs = [regions[i][0] for i in range(len(regions))]
                    future = executor.submit(process_api_request_for_page, imgs, page_num)
                    futures[future] = page_num
                
                for future in concurrent.futures.as_completed(futures):
                    page_num = futures[future]
                    try:
                        extracted_text, success, last_error = future.result()
                    except Exception as e:
                        extracted_text, success, last_error = "", False, str(e)
                    page_results[page_num] = (extracted_text, success, last_error)
            
            for page_num in range(batch_start, batch_end):
                all_regions_data = []
                regions = page_tasks[page_num]
                num_regions = len(regions)
                extracted_text, success, last_error = page_results[page_num]
                
                if crop_regions:
                    if success:
                        if out_format in ["xlsx", "csv", "json", "md", "docx"]:
                            try:
                                clean_text = extracted_text.strip()
                                md_fence = "`" * 3
                                if clean_text.startswith(md_fence + "json"): clean_text = clean_text[7:]
                                if clean_text.startswith(md_fence): clean_text = clean_text[3:]
                                if clean_text.endswith(md_fence): clean_text = clean_text[:-3]
                                clean_text = clean_text.strip()
                                
                                data = json.loads(clean_text)
                                regions_data = data.get("regions", [])
                                if not regions_data and isinstance(data, list): regions_data = data
                                
                                for region_idx in range(num_regions):
                                    crop_img_array = regions[region_idx][1]
                                    rows = regions_data[region_idx] if region_idx < len(regions_data) else []
                                    if not isinstance(rows, list): rows = [rows]
                                    
                                    page_data_to_write = []
                                    h_crop, w_crop = crop_img_array.shape[:2]
                                    clean_rows = []
                                    for r in rows:
                                        val = " ".join([str(x) for x in r]) if isinstance(r, list) else str(r)
                                        if '\n' in val:
                                            lines = [l.strip() for l in val.split('\n') if l.strip()]
                                            if all(len(l) <= 2 for l in lines): val = "".join(lines)
                                        clean_rows.append(val)
                                    if h_crop > w_crop * 1.5 and all(len(x.strip()) <= 2 for x in clean_rows if x.strip()): 
                                        page_data_to_write.append(["".join(clean_rows)])
                                    else:
                                        for val in clean_rows: page_data_to_write.append([val])
                                    all_regions_data.append(page_data_to_write)
                            except Exception as e:
                                for region_idx in range(num_regions): all_regions_data.append([[f"JSONパースエラー: {e}"]])
                        else:
                            text_blocks = extracted_text.split("===REGION_SPLIT===")
                            for region_idx in range(num_regions):
                                block_text = text_blocks[region_idx] if region_idx < len(text_blocks) else ""
                                all_regions_data.append([[line] for line in block_text.strip().split('\n')])
                    else:
                        for region_idx in range(num_regions): all_regions_data.append([[f"AI抽出失敗: {last_error}"]])
                else:
                    if success:
                        if out_format in ["xlsx", "csv", "json", "md", "docx"]:
                            try:
                                clean_text = extracted_text.strip()
                                md_fence = "`" * 3
                                if clean_text.startswith(md_fence + "json"): clean_text = clean_text[7:]
                                if clean_text.startswith(md_fence): clean_text = clean_text[3:]
                                if clean_text.endswith(md_fence): clean_text = clean_text[:-3]
                                clean_text = clean_text.strip()
                                
                                data = json.loads(clean_text)
                                header = data.get("header", [])
                                rows = data.get("rows", [])
                                if not header and not rows and isinstance(data, list):
                                    if data: header, rows = (data[0] if isinstance(data[0], list) else [str(data[0])]), data[1:]
                                safe_header = [str(x).strip() for x in header] if isinstance(header, list) else []
                                page_col_count = len(safe_header)
                                for r in rows:
                                    if isinstance(r, list) and len(r) > page_col_count: page_col_count = len(r)
                                if not safe_header: safe_header = [f"列{idx+1}" for idx in range(page_col_count)]
                                page_data_to_write = []
                                padded_header = (safe_header + [""] * page_col_count)[:page_col_count]
                                page_data_to_write.append(padded_header)
                                for row_data in rows:
                                    parsed_r = parse_row_data(row_data)
                                    safe_row_local = []
                                    for val in (parsed_r + [""] * page_col_count)[:page_col_count]:
                                        v_str = str(val)
                                        if '\n' in v_str:
                                            lines = [l.strip() for l in v_str.split('\n') if l.strip()]
                                            if len(lines) > 1 and all(len(l) <= 2 for l in lines): v_str = "".join(lines)
                                        safe_row_local.append(v_str)
                                    if any(v != "" for v in safe_row_local): page_data_to_write.append(safe_row_local)
                                all_regions_data.append(page_data_to_write)
                            except Exception as e:
                                all_regions_data.append([[f"JSONパースエラー: {e}"]])
                        else:
                            all_regions_data.append([[line] for line in extracted_text.split('\n')])
                    else:
                        all_regions_data.append([[f"AI抽出失敗: {last_error}"]])
                
                merged_data = merge_2d_arrays_horizontally(all_regions_data)
                final_data = []
                page_info_str = f"{page_num+1}/{total_pages}"
                
                if crop_regions:
                    header = ["ページ番号"] + [f"抽出範囲{idx+1}" for idx in range(num_regions)]
                    final_data.append(header)
                    for row in merged_data: final_data.append([page_info_str] + row)
                else:
                    if out_format in ["xlsx", "csv", "json", "md", "docx"]:
                        for r_idx, row in enumerate(merged_data):
                            if r_idx == 0: final_data.append(["ページ番号"] + row)
                            else: final_data.append([page_info_str] + row)
                    else:
                        for row in merged_data: final_data.append([page_info_str] + row)
                
                save_path = os.path.join(save_dir, f"{base}_Page_{str(page_num+1).zfill(digits)}_AI抽出")
                if out_format == "xlsx":
                    wb = Workbook(); ws = wb.active; ws.title = f"Page_{str(page_num+1).zfill(digits)}"
                    for r_idx, row_data in enumerate(final_data, 1):
                        for c_idx, val in enumerate(row_data, 1): ws.cell(row=r_idx, column=c_idx, value=sanitize_excel_text(val))
                    auto_adjust_excel_column_width(ws); wb.save(f"{save_path}.xlsx")
                elif out_format == "csv":
                    with open(f"{save_path}.csv", "w", encoding="utf-8-sig", newline="") as f_out: csv.writer(f_out).writerows(final_data)
                elif out_format == "txt":
                    with open(f"{save_path}.txt", "w", encoding="utf-8") as f_out:
                        for row_data in final_data: f_out.write("\t".join(row_data) + "\n")
                elif out_format == "json":
                    with open(f"{save_path}.json", "w", encoding="utf-8") as f_out: json.dump(final_data, f_out, ensure_ascii=False, indent=2)
                elif out_format == "md":
                    with open(f"{save_path}.md", "w", encoding="utf-8") as f_out:
                        if final_data:
                            f_out.write("| " + " | ".join(map(str, final_data[0])) + " |\n")
                            f_out.write("|" + "|".join(["---"] * len(final_data[0])) + "|\n")
                            for row in final_data[1:]: f_out.write("| " + " | ".join(map(lambda x: str(x).replace('\n', '<br>'), row)) + " |\n")
                elif out_format == "docx":
                    if Document is None: raise Exception("python-docx ライブラリがインストールされていません。")
                    doc_out = Document()
                    if final_data:
                        table = doc_out.add_table(rows=len(final_data), cols=max(len(r) for r in final_data))
                        table.style = 'Table Grid'
                        for r_idx, row_data in enumerate(final_data):
                            row_cells = table.rows[r_idx].cells
                            for c_idx, val in enumerate(row_data):
                                if c_idx < len(row_cells):
                                    row_cells[c_idx].text = str(val)
                    doc_out.save(f"{save_path}.docx")
        doc.close(); gc.collect()
        ui.set_determinate(total_pages, total_pages, "完了")

def aggregate_local_task(files, save_dir, options, ui):
    out_format = options.get("out_format", "xlsx")
    search_ext = "xlsx" if out_format in ["jpg", "png", "dxf", "svg", "tiff", "bmp"] else out_format
    
    if not files: raise Exception("処理対象のファイルまたはフォルダが選択されていません。")

    search_exts = ["xlsx", "xlsm", "xls"] if search_ext == "xlsx" else [search_ext]

    target_files_set = set()
    for f in files:
        if os.path.isdir(f):
            try:
                for fn in os.listdir(f):
                    if any(fn.lower().endswith(f".{ext}") for ext in search_exts) and "データ集約" not in fn and not fn.startswith("~$"):
                        target_files_set.add(os.path.abspath(os.path.join(f, fn)))
            except Exception: pass
        elif os.path.isfile(f):
            if any(f.lower().endswith(f".{ext}") for ext in search_exts) and "データ集約" not in os.path.basename(f) and not os.path.basename(f).startswith("~$"):
                target_files_set.add(os.path.abspath(f))

    target_files = sorted(list(target_files_set))
    if not target_files: raise Exception(f"選択したファイルやフォルダ内に集約可能な ({', '.join(['.' + ext for ext in search_exts])}) データが見つかりません。")

    agg_header, agg_rows, agg_texts = ["元ファイル名"], [], []
    
    def map_to_master(fname, curr_header, curr_rows):
        if not curr_header: curr_header = [f"列{i+1}" for i in range(max(1, len(curr_rows[0]) if curr_rows else 1))]
        safe_header = [str(h).strip() if h is not None else "" for h in parse_row_data(curr_header)]
        for i in range(len(safe_header)):
            if not safe_header[i]: safe_header[i] = f"列{i+1}"
            
        col_mapping = {}
        for i, h in enumerate(safe_header):
            match_idx = -1
            for m_idx, m_h in enumerate(agg_header):
                if m_idx == 0: continue
                if h == str(m_h).strip() and h != "": match_idx = m_idx; break
            
            if match_idx == -1:
                target_idx = i + 1
                if target_idx < len(agg_header):
                    match_idx = target_idx
                    if str(agg_header[target_idx]).startswith("列") and not h.startswith("列"): agg_header[target_idx] = h
                else: agg_header.append(h); match_idx = len(agg_header) - 1
            col_mapping[i] = match_idx
                
        for r in curr_rows:
            r_list = parse_row_data(r)
            row = [""] * len(agg_header)
            row[0] = fname 
            for i, val in enumerate(r_list):
                if i in col_mapping:
                    m_idx = col_mapping[i]
                    if m_idx >= len(row): row.extend([""] * (m_idx - len(row) + 1))
                    row[m_idx] = str(val).strip() if val is not None and str(val).strip() != "None" else ""
            if any(v != "" for v in row[1:]): agg_rows.append(row)

    for i, f in enumerate(target_files, 1):
        if ui.is_cancelled(): return
        ui.update_overall(i, len(target_files), f"データを集約中... ( {i} / {len(target_files)} ファイル )")
        ui.set_determinate(50, 100, f"読み込み中: {os.path.basename(f)}")
        
        ext_lower = os.path.splitext(f)[1].lower().strip('.')
        # アプリで抽出したデータの場合は接尾辞を取り除いてスッキリさせ、一般のファイルならそのままの名前を使う
        fname = os.path.basename(f)
        fname = re.sub(r'(?i)(_Page_\d+)?(_AI抽出|_Tesseract抽出|_Excel|_CSV|_Text)\.' + ext_lower + '$', '.pdf', fname)
        
        try:
            if ext_lower in ["xlsx", "xlsm"]:
                wb = openpyxl.load_workbook(f, data_only=True)
                for sheet in wb.sheetnames:
                    all_rows = list(wb[sheet].iter_rows(values_only=True))
                    # 一般のExcelファイルに対応するため、最初の有効な行（空でない行）を自動的にヘッダーとして認識する
                    valid_rows = []
                    for r in all_rows:
                        if r and any(c is not None and str(c).strip() != "" for c in r):
                            valid_rows.append(r)
                    if valid_rows:
                        if len(valid_rows) > 1: map_to_master(fname, valid_rows[0], valid_rows[1:])
                        else: map_to_master(fname, valid_rows[0], [])
                wb.close()
            elif ext_lower == "xls":
                if xlrd is None: raise Exception("xlsファイルの読み込みには 'xlrd' が必要です。")
                wb = xlrd.open_workbook(f)
                for sheet_idx in range(wb.nsheets):
                    sheet = wb.sheet_by_index(sheet_idx)
                    all_rows = [sheet.row_values(r_idx) for r_idx in range(sheet.nrows)]
                    valid_rows = []
                    for r in all_rows:
                        if r and any(c is not None and str(c).strip() != "" for c in r):
                            valid_rows.append(r)
                    if valid_rows:
                        if len(valid_rows) > 1: map_to_master(fname, valid_rows[0], valid_rows[1:])
                        else: map_to_master(fname, valid_rows[0], [])
            elif ext_lower == "csv":
                rows = []
                # 他のシステムで作成された一般的なCSV（Shift_JISなど）にも対応するためのフォールバック処理
                try:
                    with open(f, "r", encoding="utf-8-sig") as f_in: rows = list(csv.reader(f_in))
                except UnicodeDecodeError:
                    try:
                        with open(f, "r", encoding="cp932") as f_in: rows = list(csv.reader(f_in))
                    except Exception:
                        with open(f, "r", encoding="utf-8", errors="ignore") as f_in: rows = list(csv.reader(f_in))
                
                valid_rows = []
                for r in rows:
                    if r and any(c.strip() != "" for c in r):
                        valid_rows.append(r)
                if valid_rows:
                    if len(valid_rows) > 1: map_to_master(fname, valid_rows[0], valid_rows[1:])
                    else: map_to_master(fname, valid_rows[0], [])
            elif search_ext == "json":
                with open(f, "r", encoding="utf-8") as f_in:
                    try:
                        data = json.load(f_in)
                        rows = data.get("rows", []) if isinstance(data, dict) else data
                        if rows and isinstance(rows, list) and len(rows) > 0:
                            if len(rows) > 1: map_to_master(fname, rows[0], rows[1:])
                            else: map_to_master(fname, rows[0], [])
                    except Exception: pass
            elif search_ext == "md":
                with open(f, "r", encoding="utf-8") as f_in:
                    rows = []
                    for line in f_in:
                        line = line.strip()
                        if line.startswith('|') and line.endswith('|'):
                            cols = [c.strip().replace('<br>', '\n') for c in line[1:-1].split('|')]
                            if all(c.strip() == '-' * len(c.strip()) or c.strip() == '' or ':' in c for c in cols): continue
                            rows.append(cols)
                    if rows and len(rows) > 0:
                        if len(rows) > 1: map_to_master(fname, rows[0], rows[1:])
                        else: map_to_master(fname, rows[0], [])
            elif search_ext == "docx":
                if Document is None: raise Exception("python-docx ライブラリがインストールされていません。")
                doc_in = Document(f)
                for table in doc_in.tables:
                    rows = []
                    for row in table.rows:
                        rows.append([cell.text for cell in row.cells])
                    if rows and len(rows) > 0:
                        if len(rows) > 1: map_to_master(fname, rows[0], rows[1:])
                        else: map_to_master(fname, rows[0], [])
            elif search_ext == "txt":
                text_content = ""
                try:
                    with open(f, "r", encoding="utf-8-sig") as f_in: text_content = f_in.read()
                except UnicodeDecodeError:
                    try:
                        with open(f, "r", encoding="cp932") as f_in: text_content = f_in.read()
                    except Exception:
                        with open(f, "r", encoding="utf-8", errors="ignore") as f_in: text_content = f_in.read()
                if text_content.strip():
                    agg_texts.append(f"[{fname}]\n{text_content}")
        except Exception as e: 
            print(f"Read Error in {f}: {e}")
            
        ui.set_determinate(100, 100, "完了"); time.sleep(0.05)
        
    if len(target_files) > 0 and save_dir:
        if ui.is_cancelled(): return
        ui.set_indeterminate("集約データを保存中...")
        if search_ext in ["xlsx", "csv", "json", "md", "docx"]:
            final_data = [agg_header] + [(r + [""] * len(agg_header))[:len(agg_header)] for r in agg_rows]
            apply_text_inheritance(final_data)
            
            if search_ext == "xlsx" and len(final_data) > 1:
                wb = Workbook(); ws = wb.active; ws.title = "集約データ"
                for r_idx, r_data in enumerate(final_data, 1):
                    for c_idx, val in enumerate(r_data, 1): 
                        ws.cell(row=r_idx, column=c_idx, value=sanitize_excel_text(val))
                auto_adjust_excel_column_width(ws); wb.save(os.path.join(save_dir, "データ集約.xlsx"))
            elif search_ext == "csv" and len(final_data) > 1:
                with open(os.path.join(save_dir, "データ集約.csv"), "w", encoding="utf-8-sig", newline="") as f_out: 
                    csv.writer(f_out).writerows(final_data)
            elif search_ext == "json" and len(final_data) > 1:
                with open(os.path.join(save_dir, "データ集約.json"), "w", encoding="utf-8") as f_out:
                    json.dump(final_data, f_out, ensure_ascii=False, indent=2)
            elif search_ext == "md" and len(final_data) > 1:
                with open(os.path.join(save_dir, "データ集約.md"), "w", encoding="utf-8") as f_out:
                    f_out.write("| " + " | ".join(map(str, final_data[0])) + " |\n")
                    f_out.write("|" + "|".join(["---"] * len(final_data[0])) + "|\n")
                    for row in final_data[1:]: f_out.write("| " + " | ".join(map(lambda x: str(x).replace('\n', '<br>'), row)) + " |\n")
            elif search_ext == "docx" and len(final_data) > 1:
                if Document is None: raise Exception("python-docx ライブラリがインストールされていません。")
                doc_out = Document()
                table = doc_out.add_table(rows=len(final_data), cols=max(len(r) for r in final_data))
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(final_data):
                    row_cells = table.rows[r_idx].cells
                    for c_idx, val in enumerate(row_data):
                        if c_idx < len(row_cells):
                            row_cells[c_idx].text = str(val)
                doc_out.save(os.path.join(save_dir, "データ集約.docx"))
        elif search_ext == "txt" and agg_texts:
            with open(os.path.join(save_dir, "データ集約.txt"), "w", encoding="utf-8") as f_out: 
                f_out.write("\n\n".join(agg_texts))

def aggregate_gemini_task(files, save_dir, options, ui):
    out_format = options.get("out_format", "xlsx")
    search_ext = "xlsx" if out_format in ["jpg", "png", "dxf", "svg", "tiff", "bmp"] else out_format
    
    if not files: raise Exception("処理対象のファイルまたはフォルダが選択されていません。")

    search_exts = ["xlsx", "xlsm", "xls"] if search_ext == "xlsx" else [search_ext]

    target_files_set = set()
    for f in files:
        if os.path.isdir(f):
            try:
                for fn in os.listdir(f):
                    if any(fn.lower().endswith(f".{ext}") for ext in search_exts) and "データ集約" not in fn and not fn.startswith("~$"):
                        target_files_set.add(os.path.abspath(os.path.join(f, fn)))
            except Exception: pass
        elif os.path.isfile(f):
            if any(f.lower().endswith(f".{ext}") for ext in search_exts) and "データ集約" not in os.path.basename(f) and not os.path.basename(f).startswith("~$"):
                target_files_set.add(os.path.abspath(f))

    target_files = sorted(list(target_files_set))
    if not target_files: raise Exception(f"選択したファイルやフォルダ内に集約可能な ({', '.join(['.' + ext for ext in search_exts])}) データが見つかりません。")

    api_key = options.get("api_key", "")
    if not api_key: raise Exception("Gemini APIキーが設定されていません。「⚙️ 詳細設定」ボタンからAPIキーを入力してください。")
    genai.configure(api_key=api_key)
    models_to_try = options.get("models_to_try", ["gemini-2.5-flash"])
    model_name = models_to_try[0]

    all_file_contents = []
    
    for i, f in enumerate(target_files, 1):
        if ui.is_cancelled(): return
        ui.update_overall(i, len(target_files), f"データを読み込み中... ( {i} / {len(target_files)} ファイル )")
        ui.set_determinate(50, 100, f"読み込み: {os.path.basename(f)}")
        
        ext_lower = os.path.splitext(f)[1].lower().strip('.')
        fname = os.path.basename(f)
        fname = re.sub(r'(?i)(_Page_\d+)?(_AI抽出|_Tesseract抽出|_Excel|_CSV|_Text)\.' + ext_lower + '$', '.pdf', fname)
        
        file_data_str = f"--- [元ファイル名: {fname}] ---\n"
        
        try:
            if ext_lower in ["xlsx", "xlsm"]:
                wb = openpyxl.load_workbook(f, data_only=True)
                for sheet in wb.sheetnames:
                    for r in wb[sheet].iter_rows(values_only=True):
                        # 空行を無視してデータがある行だけを抽出
                        if r and any(c is not None and str(c).strip() != "" for c in r):
                            file_data_str += "\t".join([str(c).replace('\n', ' ') if c is not None else "" for c in r]) + "\n"
                wb.close()
            elif ext_lower == "xls":
                if xlrd is None: raise Exception("xlsファイルの読み込みには 'xlrd' が必要です。")
                wb = xlrd.open_workbook(f)
                for sheet_idx in range(wb.nsheets):
                    sheet = wb.sheet_by_index(sheet_idx)
                    for r_idx in range(sheet.nrows):
                        r = sheet.row_values(r_idx)
                        if r and any(c is not None and str(c).strip() != "" for c in r):
                            file_data_str += "\t".join([str(c).replace('\n', ' ') if c is not None else "" for c in r]) + "\n"
            elif ext_lower == "csv":
                rows = []
                try:
                    with open(f, "r", encoding="utf-8-sig") as f_in: rows = list(csv.reader(f_in))
                except UnicodeDecodeError:
                    try:
                        with open(f, "r", encoding="cp932") as f_in: rows = list(csv.reader(f_in))
                    except Exception:
                        with open(f, "r", encoding="utf-8", errors="ignore") as f_in: rows = list(csv.reader(f_in))
                for r in rows:
                    if any(c.strip() != "" for c in r):
                        file_data_str += "\t".join([c.replace('\n', ' ') for c in r]) + "\n"
            elif search_ext == "json":
                with open(f, "r", encoding="utf-8") as f_in:
                    try:
                        data = json.load(f_in)
                        rows = data.get("rows", []) if isinstance(data, dict) else data
                        if rows and isinstance(rows, list):
                            for r in rows:
                                if isinstance(r, list): file_data_str += "\t".join([str(c).replace('\n', ' ') for c in r]) + "\n"
                                else: file_data_str += str(r).replace('\n', ' ') + "\n"
                    except: pass
            elif search_ext == "md":
                with open(f, "r", encoding="utf-8") as f_in:
                    for line in f_in:
                        line = line.strip()
                        if line.startswith('|') and line.endswith('|'):
                            cols = [c.strip().replace('<br>', ' ') for c in line[1:-1].split('|')]
                            if all(c.strip() == '-' * len(c.strip()) or c.strip() == '' or ':' in c for c in cols): continue
                            file_data_str += "\t".join(cols) + "\n"
            elif search_ext == "docx":
                if Document is not None:
                    doc_in = Document(f)
                    for table in doc_in.tables:
                        for row in table.rows:
                            file_data_str += "\t".join([cell.text.replace('\n', ' ') for cell in row.cells]) + "\n"
            elif search_ext == "txt":
                text_content = ""
                try:
                    with open(f, "r", encoding="utf-8-sig") as f_in: text_content = f_in.read()
                except UnicodeDecodeError:
                    try:
                        with open(f, "r", encoding="cp932") as f_in: text_content = f_in.read()
                    except Exception:
                        with open(f, "r", encoding="utf-8", errors="ignore") as f_in: text_content = f_in.read()
                file_data_str += text_content + "\n"
        except Exception as e:
            print(f"Read Error in {f}: {e}")
            
        all_file_contents.append(file_data_str)
        ui.set_determinate(100, 100, "完了"); time.sleep(0.05)
        
    if not all_file_contents:
        raise Exception("集約するデータが抽出できませんでした。")

    combined_text = "\n".join(all_file_contents)
    
    if ui.is_cancelled(): return
    ui.set_indeterminate("Gemini APIでデータを統合・集約中...")

    prompt = """
    あなたは優秀なデータアナリストです。
    提供された複数のファイルの表データを解析し、意味的に同じ列を自動的に統合して、1つのマスターデータ（JSON形式）を作成してください。

    【ルール - 絶対厳守】
    1. 各行の最初の列には必ず「元ファイル名」という列を設け、そのデータがどのファイルから来たかを明記してください。
    2. 列名が異なっていても（例:「氏名」と「名前」、「単価」と「金額」など）、文脈から同じ意味の列と判断できる場合は同じ列に結合してください。
    3. JSONデータ以外は絶対に出力しないでください（マークダウンの ```json なども不要です）。
    4. 以下の形式に厳密に従ってください。
    {
      "header": ["元ファイル名", "統合列名1", "統合列名2", ...],
      "rows": [
        ["ファイルA.pdf", "データ1", "データ2", ...],
        ["ファイルB.pdf", "データ1", "", ...]
      ]
    }
    """

    generation_config = {"response_mime_type": "application/json"}
    
    max_retries = 5
    result_text = ""
    last_error = ""
    
    for attempt in range(max_retries):
        if ui.is_cancelled(): return
        try:
            if attempt > 0:
                ui.set_indeterminate(f"Gemini APIでデータを統合・集約中... (再試行 {attempt}/{max_retries-1})")
                
            model = genai.GenerativeModel(model_name)
            response = model.generate_content([prompt, combined_text], generation_config=generation_config)
            if not response.parts: raise Exception("安全フィルタ等によりブロックされました。")
            
            result_text = response.text.strip()
            break # 成功したのでループを抜ける
            
        except Exception as e:
            err_str = str(e)
            last_error = err_str
            if attempt < max_retries - 1 and ("429" in err_str or "Quota" in err_str):
                # 抽出処理の直後にAPIを叩くためバーストしやすい。引っかかった場合は長めに待機する
                m = re.search(r'retry in ([\d\.]+)s', err_str)
                if not m: m = re.search(r'seconds:\s*(\d+)', err_str)
                wait_sec = float(m.group(1)) + 2.0 if m else 15.0
                
                ui.set_indeterminate(f"API制限回避のため約{int(wait_sec)}秒待機中...")
                wait_step = 0.5
                while wait_sec > 0:
                    if ui.is_cancelled(): return
                    time.sleep(min(wait_step, wait_sec))
                    wait_sec -= wait_step
            else:
                raise Exception(f"Gemini APIによる集約に失敗しました: {last_error}\nデータ量が多すぎるか、形式エラーです。ローカル集約を使用してください。")
    
    try:
        md_fence = "`" * 3
        if result_text.startswith(md_fence + "json"): result_text = result_text[7:]
        if result_text.startswith(md_fence): result_text = result_text[3:]
        if result_text.endswith(md_fence): result_text = result_text[:-3]
        result_text = result_text.strip()
        
        data = json.loads(result_text)
        
        final_header = data.get("header", [])
        final_rows = data.get("rows", [])
        
        if not final_header or not final_rows:
            raise Exception("AIからの応答が不正な形式でした。")
            
        final_data = [final_header] + final_rows
        
    except Exception as e:
        raise Exception(f"Gemini APIによる集約結果の解析に失敗しました: {e}\nデータ量が多すぎるか、形式エラーです。ローカル集約を使用してください。")

    if ui.is_cancelled(): return
    ui.set_indeterminate("集約データを保存中...")

    if search_ext in ["xlsx", "csv", "json", "md", "docx"]:
        apply_text_inheritance(final_data)
        
        if search_ext == "xlsx":
            wb = Workbook(); ws = wb.active; ws.title = "集約データ"
            for r_idx, r_data in enumerate(final_data, 1):
                for c_idx, val in enumerate(r_data, 1): 
                    ws.cell(row=r_idx, column=c_idx, value=sanitize_excel_text(val))
            auto_adjust_excel_column_width(ws); wb.save(os.path.join(save_dir, "データ集約.xlsx"))
        elif search_ext == "csv":
            with open(os.path.join(save_dir, "データ集約.csv"), "w", encoding="utf-8-sig", newline="") as f_out: 
                csv.writer(f_out).writerows(final_data)
        elif search_ext == "json":
            with open(os.path.join(save_dir, "データ集約.json"), "w", encoding="utf-8") as f_out:
                json.dump(final_data, f_out, ensure_ascii=False, indent=2)
        elif search_ext == "md":
            with open(os.path.join(save_dir, "データ集約.md"), "w", encoding="utf-8") as f_out:
                f_out.write("| " + " | ".join(map(str, final_data[0])) + " |\n")
                f_out.write("|" + "|".join(["---"] * len(final_data[0])) + "|\n")
                for row in final_data[1:]: f_out.write("| " + " | ".join(map(lambda x: str(x).replace('\n', '<br>'), row)) + " |\n")
        elif search_ext == "docx":
            if Document is None: raise Exception("python-docx ライブラリがインストールされていません。")
            doc_out = Document()
            table = doc_out.add_table(rows=len(final_data), cols=max(len(r) for r in final_data))
            table.style = 'Table Grid'
            for r_idx, row_data in enumerate(final_data):
                row_cells = table.rows[r_idx].cells
                for c_idx, val in enumerate(row_data):
                    if c_idx < len(row_cells):
                        row_cells[c_idx].text = str(val)
            doc_out.save(os.path.join(save_dir, "データ集約.docx"))
    elif search_ext == "txt":
        with open(os.path.join(save_dir, "データ集約.txt"), "w", encoding="utf-8") as f_out: 
            f_out.write("\n\n".join(all_file_contents))