# -*- coding: utf-8 -*-
import os, re, csv, json, time, random
import openpyxl
from openpyxl import Workbook

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import xlrd
except ImportError:
    xlrd = None

from utils import (
    auto_adjust_excel_column_width,
    parse_row_data,
    apply_text_inheritance,
    sanitize_excel_text
)

# ==============================
# データ集約タスク (ローカル)
# ==============================
def aggregate_local_task(files, save_dir, options, ui):
    out_format = options.get("out_format", "xlsx")
    search_ext = "xlsx" if out_format in ["jpg", "png", "dxf", "svg", "tiff", "bmp"] else out_format
    
    if not files: raise Exception("処理対象のファイルまたはフォルダが選択されていません。")

    search_exts = ["xlsx", "xlsm", "xls"] if search_ext == "xlsx" else [search_ext]

    target_files_set = set()
    for f in files:
        if os.path.isdir(f):
            try:
                for fn in os.listdir(f):
                    if any(fn.lower().endswith(f".{ext}") for ext in search_exts) and "データ集約" not in fn and not fn.startswith("~$"):
                        target_files_set.add(os.path.abspath(os.path.join(f, fn)))
            except Exception: pass
        elif os.path.isfile(f):
            if any(f.lower().endswith(f".{ext}") for ext in search_exts) and "データ集約" not in os.path.basename(f) and not os.path.basename(f).startswith("~$"):
                target_files_set.add(os.path.abspath(f))

    target_files = sorted(list(target_files_set))
    if not target_files: raise Exception(f"選択したファイルやフォルダ内に集約可能な ({', '.join(['.' + ext for ext in search_exts])}) データが見つかりません。")

    agg_header, agg_rows, agg_texts = ["元ファイル名"], [], []
    
    def map_to_master(fname, curr_header, curr_rows):
        if not curr_header: curr_header = [f"列{i+1}" for i in range(max(1, len(curr_rows[0]) if curr_rows else 1))]
        safe_header = [str(h).strip() if h is not None else "" for h in parse_row_data(curr_header)]
        for i in range(len(safe_header)):
            if not safe_header[i]: safe_header[i] = f"列{i+1}"
            
        col_mapping = {}
        for i, h in enumerate(safe_header):
            match_idx = -1
            for m_idx, m_h in enumerate(agg_header):
                if m_idx == 0: continue
                if h == str(m_h).strip() and h != "": match_idx = m_idx; break
            
            if match_idx == -1:
                target_idx = i + 1
                if target_idx < len(agg_header):
                    match_idx = target_idx
                    if str(agg_header[target_idx]).startswith("列") and not h.startswith("列"): agg_header[target_idx] = h
                else: agg_header.append(h); match_idx = len(agg_header) - 1
            col_mapping[i] = match_idx
                
        for r in curr_rows:
            r_list = parse_row_data(r)
            row = [""] * len(agg_header)
            row[0] = fname 
            for i, val in enumerate(r_list):
                if i in col_mapping:
                    m_idx = col_mapping[i]
                    if m_idx >= len(row): row.extend([""] * (m_idx - len(row) + 1))
                    row[m_idx] = str(val).strip() if val is not None and str(val).strip() != "None" else ""
            if any(v != "" for v in row[1:]): agg_rows.append(row)

    for i, f in enumerate(target_files, 1):
        if ui.is_cancelled(): return
        ui.update_overall(i, len(target_files), f"データを集約中... ( {i} / {len(target_files)} ファイル )")
        ui.set_determinate(50, 100, f"読み込み中: {os.path.basename(f)}")
        
        ext_lower = os.path.splitext(f)[1].lower().strip('.')
        fname = os.path.basename(f)
        fname = re.sub(r'(?i)(_Page_\d+)?(_AI抽出|_Tesseract抽出|_Excel|_CSV|_Text)\.' + ext_lower + '$', '.pdf', fname)
        
        try:
            if ext_lower in ["xlsx", "xlsm"]:
                wb = openpyxl.load_workbook(f, data_only=True)
                for sheet in wb.sheetnames:
                    all_rows = list(wb[sheet].iter_rows(values_only=True))
                    valid_rows = []
                    for r in all_rows:
                        if r and any(c is not None and str(c).strip() != "" for c in r):
                            valid_rows.append(r)
                    if valid_rows:
                        if len(valid_rows) > 1: map_to_master(fname, valid_rows[0], valid_rows[1:])
                        else: map_to_master(fname, valid_rows[0], [])
                wb.close()
            elif ext_lower == "xls":
                if xlrd is None: raise Exception("xlsファイルの読み込みには 'xlrd' が必要です。")
                wb = xlrd.open_workbook(f)
                for sheet_idx in range(wb.nsheets):
                    sheet = wb.sheet_by_index(sheet_idx)
                    all_rows = [sheet.row_values(r_idx) for r_idx in range(sheet.nrows)]
                    valid_rows = []
                    for r in all_rows:
                        if r and any(c is not None and str(c).strip() != "" for c in r):
                            valid_rows.append(r)
                    if valid_rows:
                        if len(valid_rows) > 1: map_to_master(fname, valid_rows[0], valid_rows[1:])
                        else: map_to_master(fname, valid_rows[0], [])
            elif ext_lower == "csv":
                rows = []
                try:
                    with open(f, "r", encoding="utf-8-sig") as f_in: rows = list(csv.reader(f_in))
                except UnicodeDecodeError:
                    try:
                        with open(f, "r", encoding="cp932") as f_in: rows = list(csv.reader(f_in))
                    except Exception:
                        with open(f, "r", encoding="utf-8", errors="ignore") as f_in: rows = list(csv.reader(f_in))
                
                valid_rows = []
                for r in rows:
                    if r and any(c.strip() != "" for c in r):
                        valid_rows.append(r)
                if valid_rows:
                    if len(valid_rows) > 1: map_to_master(fname, valid_rows[0], valid_rows[1:])
                    else: map_to_master(fname, valid_rows[0], [])
            elif search_ext == "json":
                with open(f, "r", encoding="utf-8") as f_in:
                    try:
                        data = json.load(f_in)
                        rows = data.get("rows", []) if isinstance(data, dict) else data
                        if rows and isinstance(rows, list) and len(rows) > 0:
                            if len(rows) > 1: map_to_master(fname, rows[0], rows[1:])
                            else: map_to_master(fname, rows[0], [])
                    except Exception: pass
            elif search_ext == "md":
                with open(f, "r", encoding="utf-8") as f_in:
                    rows = []
                    for line in f_in:
                        line = line.strip()
                        if line.startswith('|') and line.endswith('|'):
                            cols = [c.strip().replace('<br>', '\n') for c in line[1:-1].split('|')]
                            if all(c.strip() == '-' * len(c.strip()) or c.strip() == '' or ':' in c for c in cols): continue
                            rows.append(cols)
                    if rows and len(rows) > 0:
                        if len(rows) > 1: map_to_master(fname, rows[0], rows[1:])
                        else: map_to_master(fname, rows[0], [])
            elif search_ext == "docx":
                if Document is None: raise Exception("python-docx ライブラリがインストールされていません。")
                doc_in = Document(f)
                for table in doc_in.tables:
                    rows = []
                    for row in table.rows:
                        rows.append([cell.text for cell in row.cells])
                    if rows and len(rows) > 0:
                        if len(rows) > 1: map_to_master(fname, rows[0], rows[1:])
                        else: map_to_master(fname, rows[0], [])
            elif search_ext == "txt":
                text_content = ""
                try:
                    with open(f, "r", encoding="utf-8-sig") as f_in: text_content = f_in.read()
                except UnicodeDecodeError:
                    try:
                        with open(f, "r", encoding="cp932") as f_in: text_content = f_in.read()
                    except Exception:
                        with open(f, "r", encoding="utf-8", errors="ignore") as f_in: text_content = f_in.read()
                if text_content.strip():
                    agg_texts.append(f"[{fname}]\n{text_content}")
        except Exception as e: 
            print(f"Read Error in {f}: {e}")
            
        ui.set_determinate(100, 100, "完了"); time.sleep(0.05)
        
    if len(target_files) > 0 and save_dir:
        if ui.is_cancelled(): return
        ui.set_indeterminate("集約データを保存中...")
        if search_ext in ["xlsx", "csv", "json", "md", "docx"]:
            final_data = [agg_header] + [(r + [""] * len(agg_header))[:len(agg_header)] for r in agg_rows]
            apply_text_inheritance(final_data)
            
            if search_ext == "xlsx" and len(final_data) > 1:
                wb = Workbook(); ws = wb.active; ws.title = "集約データ"
                for r_idx, r_data in enumerate(final_data, 1):
                    for c_idx, val in enumerate(r_data, 1): 
                        ws.cell(row=r_idx, column=c_idx, value=sanitize_excel_text(val))
                auto_adjust_excel_column_width(ws); wb.save(os.path.join(save_dir, "データ集約.xlsx"))
            elif search_ext == "csv" and len(final_data) > 1:
                with open(os.path.join(save_dir, "データ集約.csv"), "w", encoding="utf-8-sig", newline="") as f_out: 
                    csv.writer(f_out).writerows(final_data)
            elif search_ext == "json" and len(final_data) > 1:
                with open(os.path.join(save_dir, "データ集約.json"), "w", encoding="utf-8") as f_out:
                    json.dump(final_data, f_out, ensure_ascii=False, indent=2)
            elif search_ext == "md" and len(final_data) > 1:
                with open(os.path.join(save_dir, "データ集約.md"), "w", encoding="utf-8") as f_out:
                    f_out.write("| " + " | ".join(map(str, final_data[0])) + " |\n")
                    f_out.write("|" + "|".join(["---"] * len(final_data[0])) + "|\n")
                    for row in final_data[1:]: f_out.write("| " + " | ".join(map(lambda x: str(x).replace('\n', '<br>'), row)) + " |\n")
            elif search_ext == "docx" and len(final_data) > 1:
                if Document is None: raise Exception("python-docx ライブラリがインストールされていません。")
                doc_out = Document()
                table = doc_out.add_table(rows=len(final_data), cols=max(len(r) for r in final_data))
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(final_data):
                    row_cells = table.rows[r_idx].cells
                    for c_idx, val in enumerate(row_data):
                        if c_idx < len(row_cells):
                            row_cells[c_idx].text = str(val)
                doc_out.save(os.path.join(save_dir, "データ集約.docx"))
        elif search_ext == "txt" and agg_texts:
            with open(os.path.join(save_dir, "データ集約.txt"), "w", encoding="utf-8") as f_out: 
                f_out.write("\n\n".join(agg_texts))